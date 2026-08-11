"""Render a RAG analysis as a Word (.docx) memo.

This module is deliberately free of Django imports so the document-generation
logic can be unit-tested (and reasoned about) without a database or request.
The API view (`api/rag_workspace.RagReportDocxAPIView`) assembles a plain
``ReportPayload`` from the ORM and calls :func:`build_report_docx`, which
returns the ``.docx`` bytes.

The layout mirrors the answer-first web view a jurist reads: the question and
general summary lead, then the rule of law, the closed-question verdict, the
key considerations and the structured tables, and finally every cited decision
with its En fait / Subsomption / Conclusion reasoning and verbatim citation —
ready to drop into a memo.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, Sequence

from docx import Document
from docx.shared import Pt, RGBColor

# Section aliases, kept in sync with the frontend parser
# (`src/frontend/src/api/rag.ts`). The per-decision ``answer`` is a JSON string
# whose keys vary by pipeline; we normalise them to facts / subsumption /
# conclusion.
SECTION_KEYS: dict[str, tuple[str, ...]] = {
    "facts": ("Faits", "Mineure-Faits", "Examen des faits"),
    "subsumption": (
        "Subsomption",
        "Mineure-Subsomption",
        "Subsommation",
        "Mineure-Subsommation",
        "Analyse juridique (subsomption)",
    ),
    "conclusion": ("Conclusion", "Décision finale"),
}

_SECTION_LABELS: tuple[tuple[str, str], ...] = (
    ("facts", "En fait"),
    ("subsumption", "Subsomption"),
    ("conclusion", "Conclusion"),
)

_HEADING_COLOR = RGBColor(0x1E, 0x3A, 0x8A)  # matches the app's primary blue


def _as_text(value: Any) -> str:
    if isinstance(value, (list, tuple)):
        return " ".join(str(item) for item in value).strip()
    if value is None:
        return ""
    return str(value).strip()


def parse_sections(answer_text: str) -> dict[str, str] | None:
    """Parse a per-decision ``answer`` JSON string into labelled sections.

    Returns ``None`` when the text is not section-structured JSON, so the caller
    can fall back to rendering the raw answer.
    """
    import json

    try:
        parsed = json.loads(answer_text)
    except (ValueError, TypeError):
        return None
    if not isinstance(parsed, dict):
        return None

    has_any = any(
        any(key in parsed for key in keys) for keys in SECTION_KEYS.values()
    )
    if not has_any:
        return None

    result: dict[str, str] = {}
    for field, keys in SECTION_KEYS.items():
        text = ""
        for key in keys:
            if key in parsed:
                text = _as_text(parsed[key])
                if text:
                    break
        result[field] = text
    return result


def _add_meta_line(document: Any, answer: dict[str, Any]) -> str:
    parts = [
        _as_text(answer.get("procedure_type")),
        _as_text(answer.get("decision_type")),
        _as_text(answer.get("result")),
        _as_text(answer.get("decision_date")),
    ]
    return " · ".join(part for part in parts if part)


def _add_table(
    document: Any, title: str, rows: Sequence[dict[str, Any]]
) -> None:
    """Render a list of uniform dict rows as a simple table under a heading."""
    cleaned = [row for row in rows if isinstance(row, dict)]
    if not cleaned:
        return
    # Columns from the first row, dropping the link-only ``article_url`` helper.
    columns = [key for key in cleaned[0].keys() if key != "article_url"]
    if not columns:
        return

    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for index, column in enumerate(columns):
        header_cells[index].text = (
            "Références"
            if column == "references"
            else (column.replace("_", " ").capitalize())
        )
    for row in cleaned:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            value = row.get(column, "")
            if column == "references" and isinstance(value, list):
                cells[index].text = ", ".join(
                    _as_text(ref.get("procedure_type"))
                    for ref in value
                    if isinstance(ref, dict)
                )
            elif isinstance(value, (list, tuple)):
                cells[index].text = ", ".join(_as_text(item) for item in value)
            else:
                cells[index].text = _as_text(value)


def build_report_docx(payload: dict[str, Any]) -> bytes:
    """Render ``payload`` into a ``.docx`` memo and return its bytes."""
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_heading("LiteRev — Analyse juridique", level=0)
    for run in title.runs:
        run.font.color.rgb = _HEADING_COLOR

    question = _as_text(payload.get("question"))
    if question:
        document.add_heading("Question", level=1)
        document.add_paragraph(question)

    summary = _as_text(payload.get("summary"))
    if summary:
        document.add_heading("Réponse", level=1)
        document.add_paragraph(summary)

    if payload.get("show_closed_stats"):
        counts = payload.get("counts") or {}
        percentages = payload.get("percentages") or {}
        labels = {
            "oui": "Oui",
            "non": "Non",
            "peut_etre": "Peut-être",
            "mixte": "Mixte",
        }
        parts = [
            f"{label} : {counts.get(key, 0)} ({percentages.get(key, 0)}%)"
            for key, label in labels.items()
        ]
        verdict = document.add_paragraph()
        verdict.add_run("Verdict — ").bold = True
        verdict.add_run(" · ".join(parts))

    regle_droit = _as_text(payload.get("regle_droit"))
    if regle_droit:
        document.add_heading("Règle de droit", level=1)
        document.add_paragraph(regle_droit)

    considerations = payload.get("considerations") or []
    rendered = [
        c
        for c in considerations
        if isinstance(c, dict) and _as_text(c.get("text"))
    ]
    if rendered:
        document.add_heading("Considérations clés", level=1)
        for consideration in rendered:
            percent = consideration.get("percent", 0) or 0
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(_as_text(consideration.get("text")))
            paragraph.add_run(f"  ({round(float(percent))}%)").italic = True

    _add_table(document, "Éléments clés", payload.get("key_elements") or [])
    _add_table(document, "Articles de loi", payload.get("law_articles") or [])

    answers = [
        a for a in (payload.get("answers") or []) if isinstance(a, dict)
    ]
    if answers:
        document.add_heading("Décisions citées", level=1)
        for index, answer in enumerate(answers, start=1):
            meta = _add_meta_line(document, answer)
            heading = document.add_heading(
                f"[{index}] {meta}" if meta else f"[{index}]", level=2
            )
            for run in heading.runs:
                run.font.color.rgb = _HEADING_COLOR

            sections = parse_sections(_as_text(answer.get("answer")))
            if sections:
                for field, label in _SECTION_LABELS:
                    text = sections.get(field, "")
                    if text:
                        paragraph = document.add_paragraph()
                        paragraph.add_run(f"{label} : ").bold = True
                        paragraph.add_run(text)
            else:
                plain = _as_text(answer.get("answer"))
                if plain:
                    document.add_paragraph(plain)

            citation = _as_text(answer.get("citation"))
            if citation:
                quote = document.add_paragraph(f"« {citation} »")
                quote.style = "Intense Quote"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
