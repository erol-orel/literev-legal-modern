"""Unit tests for the RAG Word (.docx) memo renderer (``libs.report_docx``).

DB- and request-free: ``build_report_docx`` is a pure function over a plain
payload. The tests parse section JSON, render a full report, and re-open the
produced ``.docx`` to assert the question, rule of law, verdict and cited
decisions all made it into the document.
"""

from __future__ import annotations

import json

from io import BytesIO

from docx import Document

from literev.libs.report_docx import build_report_docx, parse_sections


def _text(payload: dict) -> str:
    document = Document(BytesIO(build_report_docx(payload)))
    paragraphs = "\n".join(p.text for p in document.paragraphs)
    cells = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return f"{paragraphs}\n{cells}"


class TestParseSections:
    def test_maps_aliased_keys(self) -> None:
        parsed = parse_sections(
            json.dumps(
                {
                    "Faits": "F",
                    "Analyse juridique (subsomption)": "S",
                    "Décision finale": "C",
                }
            )
        )
        assert parsed == {"facts": "F", "subsumption": "S", "conclusion": "C"}

    def test_returns_none_for_non_json(self) -> None:
        assert parse_sections("plain answer") is None

    def test_returns_none_when_no_section_keys(self) -> None:
        assert parse_sections(json.dumps({"other": "x"})) is None


class TestBuildReportDocx:
    def _payload(self) -> dict:
        return {
            "question": "Le bailleur peut-il résilier le bail ?",
            "summary": "Dans la majorité des cas, la résiliation est admise.",
            "regle_droit": "Art. 271 CO.",
            "show_closed_stats": True,
            "counts": {"oui": 3, "non": 1, "peut_etre": 0, "mixte": 0},
            "percentages": {"oui": 75, "non": 25, "peut_etre": 0, "mixte": 0},
            "considerations": [
                {"text": "Un motif sérieux est requis", "percent": 66}
            ],
            "key_elements": [
                {
                    "element": "Congé",
                    "references": [{"procedure_type": "Bail"}],
                }
            ],
            "law_articles": [
                {
                    "article": "Art. 271 CO",
                    "article_url": "https://fedlex.test",
                }
            ],
            "answers": [
                {
                    "procedure_type": "Bail",
                    "decision_type": "Arrêt",
                    "result": "Admis",
                    "decision_date": "2023-05-12",
                    "citation": "Le bail peut être résilié.",
                    "answer": json.dumps(
                        {
                            "Faits": "Le locataire conteste.",
                            "Subsomption": "Art. 271 CO s'applique.",
                            "Conclusion": "Résiliation valable.",
                        }
                    ),
                },
                {
                    "procedure_type": "Bail",
                    "answer": "Réponse narrative simple.",
                    "citation": "",
                },
            ],
        }

    def test_returns_docx_bytes(self) -> None:
        content = build_report_docx(self._payload())
        assert isinstance(content, bytes)
        # A .docx is a ZIP container — check the magic bytes.
        assert content[:2] == b"PK"

    def test_document_contains_the_key_content(self) -> None:
        text = _text(self._payload())
        assert "Le bailleur peut-il résilier le bail ?" in text
        assert "Dans la majorité des cas" in text
        assert "Art. 271 CO." in text
        assert "Verdict" in text
        assert "Un motif sérieux est requis" in text
        # sectioned decision reasoning + verbatim citation
        assert "Le locataire conteste." in text
        assert "Résiliation valable." in text
        assert "Le bail peut être résilié." in text
        # non-sectioned answer falls back to raw text
        assert "Réponse narrative simple." in text

    def test_renders_the_two_fact_tables(self) -> None:
        document = Document(BytesIO(build_report_docx(self._payload())))
        assert len(document.tables) == 2

    def test_minimal_payload_still_renders(self) -> None:
        content = build_report_docx({"question": "Q?"})
        assert content[:2] == b"PK"

    def test_leads_cited_decision_with_swiss_citation(self) -> None:
        # An ATF reference in the citation quote is surfaced canonically in the
        # decision heading (reusing lr_legal).
        payload = {
            "question": "Q",
            "answers": [
                {
                    "procedure_type": "Bail",
                    "citation": "Selon l'ATF 145 III 72, le congé est valable.",
                    "answer": "Le tribunal confirme.",
                }
            ],
        }
        document = Document(BytesIO(build_report_docx(payload)))
        headings = [
            p.text
            for p in document.paragraphs
            if p.style.name.startswith("Heading")
        ]
        assert any("ATF 145 III 72" in h for h in headings), headings
